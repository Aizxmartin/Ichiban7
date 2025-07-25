import pandas as pd
from docx import Document
from adjustments import apply_adjustments
import json

with open("market_adjustment_schema.json", "r") as f:
    ADJUSTMENT_SCHEMA = json.load(f)

def generate_report(df, subject_info, online_estimates):
    document = Document()
    document.add_heading("Market Valuation Report – Adjusted Comparison with Breakdown", level=1)

    document.add_paragraph(f"Subject Property")
    document.add_paragraph(f"Address: {subject_info.get('Address', 'N/A')}")
    document.add_paragraph(f"Above Grade SF: {subject_info.get('Above Grade SF', 'N/A')}")
    document.add_paragraph(f"Bedrooms: {subject_info.get('Bedrooms', 'N/A')}")
    document.add_paragraph(f"Bathrooms: {subject_info.get('Bathrooms', 'N/A')}")

    # Sanitize numeric fields
    for col in ['Close Price', 'Above Grade Finished Area', 'Concessions']:
        df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)

    # Compute Net Price and PPSF
    df["Net Price"] = df["Close Price"] - df["Concessions"]
    df["PPSF"] = df["Net Price"] / df["Above Grade Finished Area"]

    # Filter AG SF range (85% - 110%)
    try:
        agsf = int(subject_info.get("Above Grade SF", 0))
        df = df[(df["Above Grade Finished Area"] >= agsf * 0.85) & 
                (df["Above Grade Finished Area"] <= agsf * 1.10)]
    except:
        pass

    # Full address field
    df["Street Address"] = df["Street Number"].astype(str).str.strip() + " " +                            df["Street Dir Prefix"].fillna("").astype(str).str.strip() + " " +                            df["Street Name"].astype(str).str.strip() + " " +                            df["Street Suffix"].astype(str).str.strip()

    df = apply_adjustments(df, subject_info, ADJUSTMENT_SCHEMA)

    # Add comp table
    document.add_paragraph("\nAdjusted Comparable Sales:")

    table = document.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Address"
    hdr[1].text = "AG SF"
    hdr[2].text = "Close Price"
    hdr[3].text = "Concessions"
    hdr[4].text = "Adjusted Price"

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row.get("Street Address", "N/A")
        row_cells[1].text = str(row.get("Above Grade Finished Area", ""))
        row_cells[2].text = str(row.get("Close Price", ""))
        row_cells[3].text = str(row.get("Concessions", ""))
        row_cells[4].text = str(row.get("Adj Close Price", ""))

    # Valuation Summary
    document.add_paragraph("\nValuation Summary")
    avg_ppsf = df["PPSF"].mean()
    avg_online = sum(online_estimates) / len(online_estimates) if online_estimates else 0
    low_range = avg_online
    high_range = df["Adj Close Price"].min(), df["Adj Close Price"].max()

    document.add_paragraph(f"Average PPSF: ${avg_ppsf:,.2f}")
    document.add_paragraph(f"Online Estimate Average: ${avg_online:,.0f}")
    document.add_paragraph(f"Recommended Range: ${low_range:,.0f} to ${high_range[1]:,.0f}")

    from io import BytesIO
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
